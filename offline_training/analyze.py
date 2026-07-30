import os
from docx import Document

ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
}

doc = Document(r'd:\Dhaniyal\Report-Generation-main\New folder\Landscaping Monthly Report - Aramex Logistics - 52A12301- April 2026.docx')

for t_idx, tbl in enumerate(doc.tables):
    has_img = False
    output = []
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            imgs = 0
            for p in cell.paragraphs:
                for run in p.runs:
                    drawings = run.element.findall('.//w:drawing', namespaces=ns)
                    imgs += sum(1 for d in drawings for b in d.findall('.//a:blip', namespaces=ns))
            text = cell.text.strip().replace('\n', ' ')
            if imgs > 0 or text:
                output.append(f"  R{r}C{c}: {imgs} images, text: '{text[:50]}'")
                if imgs > 0:
                    has_img = True
    if has_img:
        print(f"Table {t_idx}:")
        for line in output:
            print(line)
        print("---")
