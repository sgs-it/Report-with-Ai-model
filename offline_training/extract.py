import os
import io
import json
from docx import Document
from docx.oxml.shape import CT_Inline, CT_Anchor

base_dirs = [
    r"d:\Dhaniyal\Report-Generation-main\MONTHLY REPORT APRIL 2026",
    r"d:\Dhaniyal\Report-Generation-main\MONTHLY REPORT JUNE 2026"
]

out_dir = r"d:\Dhaniyal\Report-Generation-main\offline_training\dataset"
os.makedirs(out_dir, exist_ok=True)

dataset = []
img_counter = 0

def get_images_from_paragraph(paragraph, doc):
    images = []
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    for run in paragraph.runs:
        drawing_els = run.element.findall('.//w:drawing', namespaces=ns)
        for drawing in drawing_els:
            blips = drawing.findall('.//a:blip', namespaces=ns)
            for blip in blips:
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    try:
                        part = doc.part.related_parts[rId]
                        images.append(part.blob)
                    except:
                        pass
    return images

def process_cell(cell, doc):
    global img_counter
    last_images = []
    for paragraph in cell.paragraphs:
        imgs = get_images_from_paragraph(paragraph, doc)
        if imgs:
            last_images.extend(imgs)
            text = paragraph.text.strip()
            if text and len(text) > 2 and last_images:
                for blob in last_images:
                    img_counter += 1
                    img_path = os.path.join(out_dir, f"img_{img_counter}.jpg")
                    with open(img_path, "wb") as f:
                        f.write(blob)
                    dataset.append({"image": f"img_{img_counter}.jpg", "caption": text})
                last_images = []
        else:
            text = paragraph.text.strip()
            if text and len(text) > 2 and last_images:
                for blob in last_images:
                    img_counter += 1
                    img_path = os.path.join(out_dir, f"img_{img_counter}.jpg")
                    with open(img_path, "wb") as f:
                        f.write(blob)
                    dataset.append({"image": f"img_{img_counter}.jpg", "caption": text})
                last_images = []

def process_doc(path):
    print("Processing:", path)
    doc = Document(path)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_cell(cell, doc)
                
    last_images = []
    for paragraph in doc.paragraphs:
        imgs = get_images_from_paragraph(paragraph, doc)
        if imgs:
            last_images.extend(imgs)
            text = paragraph.text.strip()
            if text and len(text) > 2 and last_images:
                for blob in last_images:
                    global img_counter
                    img_counter += 1
                    img_path = os.path.join(out_dir, f"img_{img_counter}.jpg")
                    with open(img_path, "wb") as f:
                        f.write(blob)
                    dataset.append({"image": f"img_{img_counter}.jpg", "caption": text})
                last_images = []
        else:
            text = paragraph.text.strip()
            if text and len(text) > 2 and last_images:
                for blob in last_images:
                    img_counter += 1
                    img_path = os.path.join(out_dir, f"img_{img_counter}.jpg")
                    with open(img_path, "wb") as f:
                        f.write(blob)
                    dataset.append({"image": f"img_{img_counter}.jpg", "caption": text})
                last_images = []

for base_dir in base_dirs:
    if not os.path.exists(base_dir):
        continue
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            if f.endswith(".docx") and not f.startswith("~"):
                try:
                    process_doc(os.path.join(root, f))
                except Exception as e:
                    print(f"Failed to process {f}: {e}")

with open(os.path.join(out_dir, "dataset.json"), "w", encoding='utf-8') as f:
    json.dump(dataset, f, indent=2, ensure_ascii=False)
print(f"Extracted {len(dataset)} image-caption pairs.")
